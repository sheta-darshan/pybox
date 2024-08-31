from docx import Document
from docx.shared import Inches
import pytesseract
from PIL import Image
import cv2
import os
import numpy as np

def extract_images_from_docx(docx_path, output_dir):
    doc = Document(docx_path)
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    image_paths = []
    for i, rel in enumerate(doc.part.rels.values()):
        if "image" in rel.target_ref:
            image = rel.target_part.blob
            image_path = os.path.join(output_dir, f'image_{i + 1}.png')
            with open(image_path, 'wb') as img_file:
                img_file.write(image)
            image_paths.append(image_path)
    
    return image_paths

def preprocess_image(image_path):
    # Load the image using OpenCV
    img = cv2.imread(image_path)
    
    # Convert to grayscale
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    
    # Apply thresholding
    _, thresh = cv2.threshold(gray, 150, 255, cv2.THRESH_BINARY_INV)

    # Optionally, you can apply noise reduction techniques if needed
    # noise_reduced = cv2.fastNlMeansDenoising(thresh, None, 30, 7, 21)

    # Save the processed image to be used by Tesseract
    processed_image_path = os.path.join(os.path.dirname(image_path), 'processed_' + os.path.basename(image_path))
    cv2.imwrite(processed_image_path, thresh)
    
    return processed_image_path

def ocr_image(image_path):
    # Preprocess the image before OCR
    processed_image_path = preprocess_image(image_path)
    
    # Use Tesseract to do OCR on the processed image
    text = pytesseract.image_to_string(Image.open(processed_image_path))
    return text

def create_new_doc_with_text(texts, output_docx_path):
    doc = Document()
    for text in texts:
        doc.add_paragraph(text)
    doc.save(output_docx_path)

def main(docx_path, output_dir, output_docx_path):
    # Step 1: Extract images from the Word document
    image_paths = extract_images_from_docx(docx_path, output_dir)
    
    # Step 2: Perform OCR on each image with preprocessing
    extracted_texts = [ocr_image(img_path) for img_path in image_paths]
    
    # Step 3: Create a new Word document with the extracted text
    create_new_doc_with_text(extracted_texts, output_docx_path)

if __name__ == "__main__":
    input_docx_path = "/Downloads/doc1.docx"  # Path to the input Word document
    output_image_dir = "/Downloads/extracted images"  # Directory to save extracted images
    output_docx_path = "/Downloads/output12.docx"  # Path to save the output Word document
    
    main(input_docx_path, output_image_dir, output_docx_path)
    print(f"OCR completed. The extracted text has been saved to {output_docx_path}.")